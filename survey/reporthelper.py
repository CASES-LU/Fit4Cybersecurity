import os
from django.http import HttpResponse
from django.conf import settings

from survey.models import SurveyQuestion, SurveyQuestionAnswer, SurveyUser, SurveyUserAnswer, Recommendations, \
    TranslationKey
from survey.globals import TRANSLATION_UI
from utils.radarFactory import radar_factory
import matplotlib.pyplot as plt


def getRecommendations(cuser):
    allAnswers = SurveyQuestionAnswer.objects.all().order_by('question__qindex', 'aindex')

    finalReportRecs = []

    for a in allAnswers:
        userAnswer = SurveyUserAnswer.objects.all().filter(user=cuser).filter(answer=a)[0]
        recommendation = Recommendations.objects.all().filter(forAnswer=a)

        if not recommendation.exists():
            continue

        for rec in recommendation:
            if rec.min_e_count.lower() > cuser.e_count.lower() or rec.max_e_count.lower() < cuser.e_count.lower():
                continue
            if userAnswer.uvalue > 0 and rec.answerChosen:
                finalReportRecs.append(str(rec))
            elif userAnswer.uvalue <= 0 and not rec.answerChosen:
                finalReportRecs.append(str(rec))

    return finalReportRecs


def createAndSendReport(user: SurveyUser, lang):
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.style import WD_STYLE_TYPE
    from datetime import date

    filepath = settings.BASE_DIR+"/wtemps/"

    template = filepath + "template.docx"
    doc = Document(template)

    everyQuestion = SurveyQuestion.objects.all().order_by('qindex')

    introduction = ""
    file_path = os.path.join(filepath, lang.lower() + '_intro.txt')
    try:
        with open(file_path, 'r') as f:
            introduction = f.read()
    except Exception as e:
        #raise e
        raise Exception('Missing file: {}'.format(file_path))

    introduction = introduction.replace("\n\r", "\n")
    introduction = introduction.split("\n\n")
    x = 0
    for i in introduction:
        if x == 0:
            doc.add_heading(i, level=1)
            x += 1
            continue
        doc.add_paragraph(i)

    methodDescr = ""
    file_path = os.path.join(filepath, lang.lower() + '_description.txt')
    try:
        with open(file_path, 'r') as f:
            methodDescr = f.read()
    except:
        raise Exception('Missing file: {}'.format(file_path))

    methodDescr = methodDescr.replace("\n\r", "\n")
    methodDescr = methodDescr.split("\n\n")
    x = 0
    for i in methodDescr:
        if x == 0:
            doc.add_heading(i, level=1)
            x += 1
            continue
        doc.add_paragraph(i)

    results = ""
    file_path = os.path.join(filepath, lang.lower() + '_resultdisclaimer.txt')
    try:
        with open(file_path, 'r') as f:
            results = f.read()
    except:
        raise Exception('Missing file: {}'.format(file_path))

    score, detail_max, details, section_list = calculateResult(user)

    results = results.replace("\n\r", "\n")
    results = results.replace("$$result$$", str(score))
    results = results.split("\n\n")

    x = 0
    for i in results:
        if x == 0:
            doc.add_heading(i, level=1)
            x += 1

            continue
        doc.add_paragraph(i)

    chart_png_file = generate_chart_png(user, detail_max, details, section_list, lang)
    doc.add_paragraph()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(chart_png_file)


    recommendationList = getRecommendations(user)
    #recommendationList = "\n\n".join(recommendationList)

    for rec in recommendationList:
        doc.add_paragraph(rec, style='List Paragraph')



    doc.add_heading(TRANSLATION_UI['document']['questions'][lang.lower()], level=1)

    x = 1
    for i in everyQuestion:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = str(x)
        hdr_cells[1].text = str(i)

        bX = hdr_cells[0].paragraphs[0].runs[0]
        bX.font.bold = True
        bX.font.size = Pt(13)
        bX = hdr_cells[1].paragraphs[0].runs[0]
        bX.font.bold = True
        bX.font.size = Pt(13)

        answerlist = SurveyQuestionAnswer.objects.filter(question=i).order_by('aindex')

        for a in answerlist:
            row_cells = table.add_row().cells
            u = SurveyUserAnswer.objects.filter(answer=a)[0]

            if u.uvalue > 0:
                row_cells[0].text = "X"
                bX = row_cells[0].paragraphs[0].runs[0]
                bX.font.bold = True
            else:
                row_cells[0].text = " "

            row_cells[1].text = str(a)

        col = table.columns[0]
        col.width = Cm(1.5)
        col = table.columns[1]
        col.width = Cm(14.0)
        for cell in table.columns[0].cells:
            cell.width = Cm(1.5)

        doc.add_paragraph()
        x += 1

    '''
    sectorName = str(user.sector)
    # SECTOR_CHOICES is removed!
    for a,b in SECTOR_CHOICES:
        if user.sector == a:
            sectorName = str(b)

    compSize = str(user.e_count)
    for a,b in COMPANY_SIZE:
        if user.e_count == a:
            compSize = b

    recommendationList = getRecommendations(user)
    recommendationList = "\n\n".join(recommendationList)

    table = []
    ind = 0
    for i in everyQuestion:
        ind += 1
        if ind > 1:
            table.append({'ca':"", 'surveyAnswers':""})

        answerlist = SurveyQuestionAnswer.objects.filter(question=i).order_by('aindex')
        headingLine = {'ca':str(ind), 'surveyAnswers':str(i)}
        table.append(headingLine)

        for a in answerlist:
            line = {'ca':"", 'surveyAnswers':""}
            u = SurveyUserAnswer.objects.filter(answer=a)[0]

            if u.uvalue > 0:
                line['ca'] = "X"
            else:
                line['ca'] = " "

            line['surveyAnswers'] = str(a)
            table.append(line)

    everyQuestionAndAnswer = table

    document.merge(
        result=str(theResult)+"/100",
        companysize=compSize,
        resultGraph=theImage,
        #surveyAnswers=everyQuestionAndAnswer,
        ca=everyQuestionAndAnswer,
        sector=sectorName,
        generationDate=str(date.today()),
        recommendationsList=recommendationList,
        )
    '''

    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = str(date.today())+"\t\tFit4Cybersecurity"
    paragraph.style = doc.styles["Header"]


    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=result-'+lang.lower()+'.docx'
    doc.save(response)

    # make checkboxes to recommendation and a single button of get companies
    # then call getcompanies when button is hit

    return response


def calculateResult(user: SurveyUser):
    allUserAnswers = SurveyUserAnswer.objects.filter(uvalue__gt=0, user=user).order_by('answer__question__qindex',
                                                                                       'answer__aindex')
    totalscore = sum([x.answer.score for x in allUserAnswers])

    maxeval = {}
    evaluation = {}
    sectionlist = {}
    maxscore = 0

    translations = TranslationKey.objects.filter(lang=user.chosenLang, ttype='S')
    translation_key_values = {}
    for translation in translations:
        translation_key_values[translation.key] = translation.text

    for q in SurveyQuestion.objects.all():
        maxscore += q.maxPoints
        if q.section.id not in evaluation:
            evaluation[q.section.id] = 0
        if q.section.id not in maxeval:
            maxeval[q.section.id] = 0
        sectionlist[q.section.id] = translation_key_values[q.section.sectionTitleKey]

        maxeval[q.section.id] += q.maxPoints

        uanswers = SurveyUserAnswer.objects.filter(user=user, uvalue__gt=0, answer__question__id=q.id)
        scores = [x.answer.score for x in uanswers]
        evaluation[q.section.id] += sum(scores)

    # get the score in percent! with then 100 being maxscore
    totalscore = round((totalscore * 100) / maxscore)

    sectionlist = [sectionlist[x] for x in sectionlist]
    evaluation = [evaluation[x] for x in evaluation]
    maxeval = [maxeval[x] for x in maxeval]

    return totalscore, maxeval, evaluation, sectionlist


def generate_chart_png(user: SurveyUser, max_eval, evaluation, sections_list, lang):
    n = len(sections_list)
    theta = radar_factory(n, frame='polygon')

    spoke_labels = []
    for section in sections_list:
        spoke_labels.append(section)

    fig, ax = plt.subplots(figsize=(7, 5), dpi=150, nrows=1, ncols=1,
                             subplot_kw=dict(projection='radar'))
    fig.subplots_adjust(wspace=0.25, hspace=0.20, top=0.85, bottom=0.05)

    grid_step = max(max_eval) / 5
    ax.set_rgrids([0, grid_step, grid_step * 2, grid_step * 3, grid_step * 4])
    ax.set_ylim(0, max(max_eval))

    ax.plot(theta, evaluation, color='r')
    ax.fill(theta, evaluation, facecolor='r', alpha=0.25)

    ax.plot(theta, max_eval, color='b')
    ax.fill(theta, max_eval, facecolor='b', alpha=0.25)

    ax.set_varlabels(spoke_labels)

    ax.legend([TRANSLATION_UI['report']['result'][lang.lower()],
               TRANSLATION_UI['report']['resultMax'][lang.lower()]], loc=(0.9, .95),
              labelspacing=0.1, fontsize='small')

    fig.text(1.0, 1.0, TRANSLATION_UI['report']['chart'][lang.lower()],
             horizontalalignment='center', color='black', weight='bold',
             size='large')

    file_name = './static/users/survey-' + str(user.user_id) + '.png'
    plt.savefig(file_name)

    return file_name